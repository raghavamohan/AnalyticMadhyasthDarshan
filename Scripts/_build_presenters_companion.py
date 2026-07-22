#!/usr/bin/env python3
"""Build a Presenter's Companion DOCX (and optional PDF / PPTX notes sync).

Markdown source of truth convention under Studies/ or Applications/:

  Presenters-Companion-<Name>.md
  Presenters-Companion-<Name>.notes.json   # optional; 1-based slide → note text
  Presenters-Companion-<Name>.docx         # generated
  Presenters-Companion-<Name>.pdf          # generated when --pdf

Heading conventions in the companion markdown:

  # PRESENTER'S COMPANION
  # Slide N
  # <Slide title>
  ## Delivering the slide
  ## Primary-text background
  ## Likely questions from the audience

Examples (from repo root):

  python Scripts/_build_presenters_companion.py Studies/The-Ontology-of-Coexistence/Presenters-Companion-Ontology-of-Existence.md --pdf
  python Scripts/_build_presenters_companion.py Studies/.../Presenters-Companion-....md --pdf --pptx Studies/.../Deck.pptx
  python Scripts/_build_presenters_companion.py Studies/.../Presenters-Companion-....md --pptx Studies/.../Deck.pptx --notes Studies/.../Presenters-Companion-....notes.json
"""
from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from _docx_to_pdf import convert_docx_to_pdf
from _sync_pptx_speaker_notes import load_notes, sync_speaker_notes


def add_runs(paragraph, raw: str) -> None:
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", raw)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def build_docx(md: Path, docx: Path) -> int:
    lines = md.read_text(encoding="utf-8").splitlines()
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    slide_count = 0
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.startswith("# Slide "):
            slide_heading = line[2:].strip()
            title_line = ""
            if (
                i + 1 < len(lines)
                and lines[i + 1].startswith("# ")
                and not lines[i + 1].startswith("# Slide")
            ):
                title_line = lines[i + 1][2:].strip()
                i += 1
            doc.add_heading(slide_heading, level=1)
            if title_line:
                doc.add_heading(title_line, level=1)
            slide_count += 1
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            p = doc.add_paragraph()
            run = p.add_run(line[2:].strip())
            run.bold = True
            run.font.size = Pt(18)
            if "PRESENTER" in line.upper():
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.strip():
            p = doc.add_paragraph()
            add_runs(p, line.strip())
        i += 1

    docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx)
    return slide_count


def default_notes_path(md: Path) -> Path:
    return md.with_suffix(".notes.json")


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(
        description=(
            "Build a Presenter's Companion DOCX from markdown; optionally "
            "export PDF and sync PPTX speaker notes."
        )
    )
    parser.add_argument(
        "markdown",
        type=Path,
        help="Path to Presenters-Companion-*.md",
    )
    parser.add_argument(
        "--docx",
        type=Path,
        help="Output .docx path (default: same stem as the markdown)",
    )
    parser.add_argument(
        "--pdf",
        action="store_true",
        help="Also export a PDF beside the DOCX via Word COM",
    )
    parser.add_argument(
        "--pdf-output",
        type=Path,
        help="PDF output path (implies --pdf; default: same stem as the DOCX)",
    )
    parser.add_argument(
        "--pptx",
        type=Path,
        help="Optional deck whose speaker notes should be synced from JSON",
    )
    parser.add_argument(
        "--notes",
        type=Path,
        help=(
            "Notes JSON map (slide number → text). Default when --pptx is set: "
            "<markdown-stem>.notes.json"
        ),
    )
    args = parser.parse_args(argv)

    md = args.markdown.expanduser().resolve()
    if not md.is_file():
        raise SystemExit(f"Markdown not found: {md}")
    if md.suffix.lower() != ".md":
        raise SystemExit(f"Expected a .md file: {md}")

    docx = (args.docx or md.with_suffix(".docx")).expanduser().resolve()
    slide_count = build_docx(md, docx)
    print(f"Wrote {docx} ({slide_count} slides)")

    if args.pdf or args.pdf_output is not None:
        pdf = (args.pdf_output or docx.with_suffix(".pdf")).expanduser().resolve()
        convert_docx_to_pdf(docx, pdf)
        print(f"Wrote {pdf} ({pdf.stat().st_size} bytes)")

    if args.pptx is not None:
        notes_path = (
            args.notes.expanduser().resolve()
            if args.notes is not None
            else default_notes_path(md)
        )
        if not notes_path.is_file():
            raise SystemExit(
                f"--pptx requires notes JSON, but none found at {notes_path}. "
                "Pass --notes <path> or create <companion>.notes.json."
            )
        notes = load_notes(notes_path)
        sync_speaker_notes(args.pptx, notes)
        print(
            f"Updated speaker notes for {len(notes)} slides in "
            f"{args.pptx.expanduser().resolve()}"
        )

    return 0


if __name__ == "__main__":
    sys.exit(main())
